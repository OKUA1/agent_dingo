from agent_dingo.rag.base import BaseReader as _BaseReader, Document
from typing import List
import docx  # python-docx


class WordDocumentReader(_BaseReader):
    def read(self, file_path: str) -> List[Document]:
        docs = []
        doc = docx.Document(file_path)
        for i, para in enumerate(doc.paragraphs):
            text = para.text
            docs.append(Document(text, {"source": file_path, "paragraph": i}))
        return docs
